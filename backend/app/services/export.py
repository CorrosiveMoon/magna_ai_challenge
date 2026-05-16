import io
import httpx
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.enums import TA_LEFT
from docx import Document
from docx.shared import Inches, Pt


async def _fetch_image_bytes(url: str) -> bytes | None:
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(url)
            r.raise_for_status()
            return r.content
    except Exception:
        return None


async def export_pdf(title: str, body: str, image_url: str | None, created_at: datetime) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(title or "Untitled", styles["Title"]))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(created_at.strftime("%B %d, %Y"), styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    if image_url:
        img_bytes = await _fetch_image_bytes(image_url)
        if img_bytes:
            img_buf = io.BytesIO(img_bytes)
            story.append(RLImage(img_buf, width=12*cm, height=12*cm))
            story.append(Spacer(1, 0.5*cm))

    body_style = ParagraphStyle("body", parent=styles["Normal"], leading=16, spaceAfter=8, alignment=TA_LEFT)
    for para in body.split("\n\n"):
        clean = para.strip().lstrip("#").strip()
        if clean:
            story.append(Paragraph(clean, body_style))

    doc.build(story)
    return buf.getvalue()


async def export_docx(title: str, body: str, image_url: str | None, created_at: datetime) -> bytes:
    doc = Document()

    doc.add_heading(title or "Untitled", level=0)
    doc.add_paragraph(created_at.strftime("%B %d, %Y"))

    if image_url:
        img_bytes = await _fetch_image_bytes(image_url)
        if img_bytes:
            img_buf = io.BytesIO(img_bytes)
            doc.add_picture(img_buf, width=Inches(4))

    doc.add_paragraph("")
    for para in body.split("\n\n"):
        clean = para.strip()
        if clean.startswith("#"):
            level = len(clean) - len(clean.lstrip("#"))
            doc.add_heading(clean.lstrip("# ").strip(), level=min(level, 4))
        elif clean:
            p = doc.add_paragraph(clean)
            p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
